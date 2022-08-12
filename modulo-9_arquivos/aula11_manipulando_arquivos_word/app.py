import os
from docx import Document
from docx.shared import Cm

os.chdir('modulo-9_arquivos')
os.chdir('aula11_manipulando_arquivos_word')

documento = Document()
documento.add_heading('Título do Documento', 0)

paragrafo = documento.add_paragraph('Um parágrafo simples')
paragrafo.add_run(' e importante').bold = True
paragrafo.add_run(' com palavras em ')
paragrafo.add_run('itálico').italic = True

documento.add_heading('Título - Nível 1', level=1)
documento.add_heading('Título - Nível 2', level=2)
documento.add_heading('Título - Nível 3', level=3)

documento.add_paragraph('Formatação "No Spacing"', style='No Spacing')
documento.add_paragraph('Formatação "Heading 1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 1"', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_paragraph('Formatação "Title"', style='Title')
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle')
documento.add_paragraph('Formatação "Quote"', style='Quote')
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph')
documento.add_paragraph('Formatação "List Bullet"', style='List Bullet')
documento.add_paragraph('Formatação "List Number"', style='List Number')

documento.add_picture('notebook.jpg', width=Cm(5.25))

tabela = documento.add_table(rows=3, cols=2)

celula1 = tabela.cell(0, 0)
celula1.text = 'Nome'
celula2 = tabela.cell(0, 1)
celula2.text = 'Idade'

celula3 = tabela.cell(1, 0)
celula3.text = 'Rafael'
celula4 = tabela.cell(1, 1)
celula4.text = '45'

celula5 = tabela.cell(2, 0)
celula5.text = 'Amanda'
celula6 = tabela.cell(2, 1)
celula6.text = '25'

# Criando tabela através de tupla
compras = (
    (3, '101', 'Uva'),
    (7, '422', 'Pão'),
    (4, '423', 'Banana, Ovos, Tomate, Carne')
)

tabela_supermercado = documento.add_table(rows=1, cols=3)
cabecalho_tabela_supermercado = tabela_supermercado.rows[0].cells
cabecalho_tabela_supermercado[0].text = 'Quantidade'
cabecalho_tabela_supermercado[1].text = 'Id'
cabecalho_tabela_supermercado[2].text = 'Descrição'

for quantidade, id, descricao in compras:
    linha = tabela_supermercado.add_row().cells
    linha[0].text = str(quantidade)
    linha[1].text = id
    linha[2].text = descricao


documento.add_page_break() # Insere uma nova página
documento.add_paragraph('Estamos em outra página')
documento.save('demo.docx')
